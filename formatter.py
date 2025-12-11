import os
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.units import mm
import logging

logger = logging.getLogger(__name__)

class DocumentFormatter:
    # Константы для расчета объема
    CHARS_PER_PAGE = 2000  # символов на страницу А4
    WORDS_PER_PAGE = 350   # слов на страницу
    
    @staticmethod
    def calculate_volume(pages):
        """Рассчитать объем текста"""
        return {
            'chars': pages * DocumentFormatter.CHARS_PER_PAGE,
            'words': pages * DocumentFormatter.WORDS_PER_PAGE,
            'pages': pages
        }
    
    @staticmethod
    def format_for_a4(text, target_pages):
        """Форматировать текст под нужное количество страниц"""
        if not text:
            return "Информация не найдена."
        
        target_chars = target_pages * DocumentFormatter.CHARS_PER_PAGE
        
        if len(text) <= target_chars:
            return text
        
        # Обрезаем по предложениям
        sentences = []
        current_sentence = ""
        
        for char in text:
            current_sentence += char
            if char in '.!?':
                sentences.append(current_sentence.strip())
                current_sentence = ""
        
        if current_sentence:
            sentences.append(current_sentence.strip())
        
        # Собираем текст до нужного объема
        result = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence) + 1  # +1 для пробела
            
            if current_length + sentence_length <= target_chars:
                result.append(sentence)
                current_length += sentence_length
            else:
                # Если осталось место для части предложения
                remaining = target_chars - current_length - 3  # -3 для "..."
                if remaining > 20:
                    result.append(sentence[:remaining] + "...")
                break
        
        formatted_text = ' '.join(result)
        if formatted_text and not formatted_text.endswith(('.', '!', '?', '...')):
            formatted_text += '.'
        
        return formatted_text
    
    @staticmethod
    def create_word_document(content, title, work_type):
        """Создать документ Word"""
        try:
            doc = Document()
            
            # Настройка стилей
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            
            # Заголовок
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{work_type.upper()}: {title}")
            title_run.bold = True
            title_run.font.size = Pt(14)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Информация о документе
            info = doc.add_paragraph()
            info.add_run(f"Тип работы: {work_type}\n")
            info.add_run(f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n")
            info.add_run(f"Объем: {len(content)} символов\n")
            info.add_run("\n" + "="*50 + "\n\n")
            
            # Основной текст
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para)
            
            # Сохраняем во временный файл
            temp_file = tempfile.NamedTemporaryFile(
                suffix='.docx', 
                delete=False,
                prefix='konspekt_'
            )
            temp_file.close()
            
            doc.save(temp_file.name)
            logger.info(f"Создан DOCX: {temp_file.name}")
            return temp_file.name
            
        except Exception as e:
            logger.error(f"Ошибка создания DOCX: {e}")
            return None
    
    @staticmethod
    def create_pdf_document(content, title, work_type):
        """Создать PDF документ"""
        try:
            # Временный файл
            temp_file = tempfile.NamedTemporaryFile(
                suffix='.pdf',
                delete=False,
                prefix='konspekt_'
            )
            temp_file.close()
            
            # Стили
            styles = getSampleStyleSheet()
            
            # Создаем свои стили
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                alignment=TA_CENTER,
                spaceAfter=20
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                alignment=TA_LEFT,
                spaceAfter=6
            )
            
            # Документ
            doc = SimpleDocTemplate(
                temp_file.name,
                pagesize=A4,
                leftMargin=20*mm,
                rightMargin=20*mm,
                topMargin=20*mm,
                bottomMargin=20*mm
            )
            
            story = []
            
            # Заголовок
            story.append(Paragraph(f"{work_type.upper()}: {title}", title_style))
            story.append(Spacer(1, 10))
            
            # Информация
            info_text = f"""
            <b>Тип работы:</b> {work_type}<br/>
            <b>Дата:</b> {datetime.now().strftime('%d.%m.%Y %H:%M')}<br/>
            <b>Объем:</b> {len(content)} символов<br/>
            <br/>
            """
            story.append(Paragraph(info_text, normal_style))
            story.append(Spacer(1, 10))
            
            # Основной текст
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.replace('\n', '<br/>'), normal_style))
                    story.append(Spacer(1, 6))
            
            # Генерация
            doc.build(story)
            logger.info(f"Создан PDF: {temp_file.name}")
            return temp_file.name
            
        except Exception as e:
            logger.error(f"Ошибка создания PDF: {e}")
            return None
    
    @staticmethod
    def create_txt_file(content, title, work_type):
        """Создать текстовый файл"""
        try:
            temp_file = tempfile.NamedTemporaryFile(
                suffix='.txt',
                delete=False,
                prefix='konspekt_',
                mode='w',
                encoding='utf-8'
            )
            
            header = f"""
{'='*60}
{work_type.upper()}: {title}
{'='*60}
Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}
Тип: {work_type}
Объем: {len(content)} символов
{'='*60}

"""
            
            temp_file.write(header + content)
            temp_file.close()
            
            logger.info(f"Создан TXT: {temp_file.name}")
            return temp_file.name
            
        except Exception as e:
            logger.error(f"Ошибка создания TXT: {e}")
            return None
    
    @staticmethod
    def cleanup_files():
        """Очистка временных файлов"""
        import glob
        import time
        
        # Удаляем старые временные файлы
        patterns = ['konspekt_*.docx', 'konspekt_*.pdf', 'konspekt_*.txt']
        
        for pattern in patterns:
            for file in glob.glob(pattern):
                try:
                    # Удаляем файлы старше 1 часа
                    if os.path.exists(file) and time.time() - os.path.getmtime(file) > 3600:
                        os.remove(file)
                        logger.info(f"Удален старый файл: {file}")
                except Exception as e:
                    logger.error(f"Ошибка удаления {file}: {e}")
